#!/usr/bin/env python3
"""Generate a repository DOCX copy of the SafeContracts master plan.

The authoritative editable source in Git is docs/MASTER_PLAN.md. This script produces a
portable Word document during plan sync so business stakeholders can download it directly.
"""

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "MASTER_PLAN.md"
OUTPUT = ROOT / "docs" / "SafeContracts_Full_Plan_v1.2.docx"


def strip_md(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[(.*?)\]\([^)]*\)", r"\1", text)
    return text.strip()


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SafeContracts")
    run.bold = True
    run.font.size = Pt(24)
    subtitle = doc.add_paragraph("Master Product & Functional Plan — V1.2")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    in_code = False
    table_buffer = []

    def flush_table():
        nonlocal table_buffer
        if len(table_buffer) < 2:
            for row in table_buffer:
                doc.add_paragraph(strip_md(row))
            table_buffer = []
            return
        parsed = []
        for line in table_buffer:
            cells = [strip_md(c) for c in line.strip().strip("|").split("|")]
            if cells and all(re.fullmatch(r"[:\- ]+", c or "-") for c in cells):
                continue
            parsed.append(cells)
        if parsed:
            cols = max(len(r) for r in parsed)
            table = doc.add_table(rows=len(parsed), cols=cols)
            table.style = "Table Grid"
            for r_idx, row in enumerate(parsed):
                for c_idx, value in enumerate(row):
                    table.cell(r_idx, c_idx).text = value
        table_buffer = []

    for raw in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            flush_table()
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph(style="No Spacing")
            p.add_run(line).font.name = "Courier New"
            continue
        if line.startswith("|") and line.endswith("|"):
            table_buffer.append(line)
            continue
        flush_table()
        if not line.strip():
            continue
        if line.startswith("# "):
            # Main markdown title is represented by the document cover.
            continue
        if line.startswith("## "):
            doc.add_heading(strip_md(line[3:]), level=1)
        elif line.startswith("### "):
            doc.add_heading(strip_md(line[4:]), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(strip_md(line[2:]), style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(strip_md(re.sub(r"^\d+\. ", "", line)), style="List Number")
        elif line.startswith("> "):
            p = doc.add_paragraph(strip_md(line[2:]))
            p.italic = True
        else:
            doc.add_paragraph(strip_md(line))

    flush_table()
    doc.add_paragraph("Generated from docs/MASTER_PLAN.md by SafeContracts plan automation.")
    doc.save(OUTPUT)
    print(f"Generated {OUTPUT}")


if __name__ == "__main__":
    main()
