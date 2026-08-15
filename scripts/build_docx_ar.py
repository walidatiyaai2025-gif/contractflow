#!/usr/bin/env python3
"""Generate the Arabic SafeContracts V1.2 Word document from the GitHub Markdown source."""

from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "SAFE_CONTRACTS_PLAN_AR.md"
OUTPUT = ROOT / "docs" / "SafeContracts_Full_Plan_AR_v1.2.docx"


def clean(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[(.*?)\]\([^)]*\)", r"\1", text)
    return text.strip()


def rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    p = doc.add_paragraph()
    rtl(p)
    run = p.add_run("SafeContracts")
    run.bold = True
    run.font.size = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph("وثيقة الخطة الوظيفية والتنفيذية — V1.2")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table_lines = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        parsed = []
        for line in table_lines:
            cells = [clean(c) for c in line.strip().strip("|").split("|")]
            if cells and all(re.fullmatch(r"[:\- ]+", c or "-") for c in cells):
                continue
            parsed.append(cells)
        if parsed:
            cols = max(len(row) for row in parsed)
            table = doc.add_table(rows=len(parsed), cols=cols)
            table.style = "Table Grid"
            for r_i, row in enumerate(parsed):
                for c_i, value in enumerate(row):
                    cell = table.cell(r_i, c_i)
                    cell.text = value
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    for para in cell.paragraphs:
                        rtl(para)
                        if r_i == 0:
                            for rr in para.runs:
                                rr.bold = True
        table_lines = []

    for raw in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            continue
        flush_table()
        if not line.strip() or line.startswith("# "):
            continue
        if line.startswith("## "):
            p = doc.add_heading(clean(line[3:]), level=1)
            rtl(p)
        elif line.startswith("### "):
            p = doc.add_heading(clean(line[4:]), level=2)
            rtl(p)
        elif line.startswith("- "):
            p = doc.add_paragraph(clean(line[2:]), style="List Bullet")
            rtl(p)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(clean(re.sub(r"^\d+\. ", "", line)), style="List Number")
            rtl(p)
        else:
            p = doc.add_paragraph(clean(line))
            rtl(p)

    flush_table()
    p = doc.add_paragraph("تم توليد هذه النسخة تلقائيًا من docs/SAFE_CONTRACTS_PLAN_AR.md داخل GitHub.")
    rtl(p)
    doc.save(OUTPUT)
    print(f"Generated {OUTPUT}")


if __name__ == "__main__":
    main()
